import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import time
import os
import sys
import google.generativeai as genai
import logging

# --- Setup basic logging to a file and the console ---
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    handlers=[
                        logging.FileHandler("question_paper_analysis.log"),
                        logging.StreamHandler()
                    ])

def configure_gemini_api():
    """
    Configures the Gemini API. It first checks for an environment variable
    and then prompts the user in the console if it's not found.
    """
    try:
        api_key = os.environ.get("GOOGLE_API_KEY")
        if not api_key:
            logging.warning("GOOGLE_API_KEY environment variable not found.")
            # --- REQUIREMENT 1: Ask for API key in the run console ---
            api_key = input("Please enter your Google Gemini API key and press Enter: ")
        
        if not api_key:
            logging.error("API Key was not provided. Cannot proceed.")
            return None

        genai.configure(api_key=api_key)
        logging.info("Gemini API configured successfully.")
        # Use the gemini-1.5-pro model as requested in the base file structure
        return genai.GenerativeModel('gemini-2.5-pro')

    except Exception as e:
        logging.error(f"Error configuring Gemini API: {e}")
        messagebox.showerror("API Error", f"Failed to configure the Gemini API: {e}")
        return None

def read_docx_content(file_path: str) -> str:
    """Reads all text content from a .docx file."""
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        logging.error(f"Failed to read DOCX file at {file_path}: {e}")
        return None

def deduplicate_content(content: str) -> str:
    """Removes duplicate paragraphs/sections from the provided text content."""
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    unique_lines = list(dict.fromkeys(lines))
    return "\n".join(unique_lines)

def analyze_with_gemini(model, text: str) -> str:
    """
    Builds a prompt, sends it to the Gemini API for analysis based on
    Bloom's Taxonomy, and returns the result.
    """
    prompt = f"""
    Based on Bloom's Taxonomy, analyze the following question paper.
    Identify each question number, its marks, and classify it into the correct
    Bloom's Taxonomy Level (Remember, Understand, Apply, Analyze, Evaluate, Create).

    Question Paper Text:
    ---
    {text}
    ---

    Provide the output ONLY in a pipe-separated format without headers or extra text.
    Each line must represent one question.
    Format: Question_No|Marks|Blooms_Taxonomy_Level

    Example Output:
    1i)|1|Remember
    2a)|9|Apply
    """
    
    # Always sleep for 121 sec before any and every api call
    #logging.info("Pausing for 121 seconds before the API call as required...")
    #time.sleep(121)

    try:
        logging.info("Sending request to the Gemini API for analysis...")
        response = model.generate_content(prompt)
        logging.info("Successfully received a response from the API.")
        return response.text
    except Exception as e:
        logging.error(f"An error occurred during the API call: {e}")
        messagebox.showerror("API Call Failed", f"An error occurred while communicating with the Gemini API: {e}")
        return None

def parse_api_response(result_text: str) -> list:
    """Parses the pipe-separated string from the API into a list of lists."""
    parsed_data = []
    if not result_text:
        return parsed_data
    
    for line in result_text.strip().split('\n'):
        parts = [part.strip() for part in line.split('|')]
        if len(parts) == 3:
            parsed_data.append(parts)
    return parsed_data

def create_analysis_document(data: list, original_input_path: str):
    """
    Creates and saves a Word document with the analysis table in the same
    directory as the script.
    """
    if not data:
        logging.warning("No data was parsed. Cannot create analysis document.")
        return

    doc = Document()
    doc.add_heading("Bloom's Taxonomy Analysis of Question Paper", level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Question No.'
    hdr_cells[1].text = 'Marks'
    hdr_cells[2].text = 'Blooms Taxonomy Level'

    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]

    # --- REQUIREMENT 2: Save the file in the script's location ---
    # Get the base name of the original input file to create a unique output name
    base_name = os.path.basename(original_input_path)
    file_name_without_ext = os.path.splitext(base_name)[0]
    output_filename = f"{file_name_without_ext}-analysis.docx"
    
    # Determine the directory where the script is located
    # sys.argv[0] is the script path.
    script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
    save_path = os.path.join(script_dir, output_filename)
    
    doc.save(save_path)
    logging.info(f"Analysis document saved successfully at: {save_path}")
    messagebox.showinfo("Success", f"Analysis complete! The report has been saved in the script's directory as:\n{save_path}")

def main():
    """Main function to drive the script."""
    logging.info("--- Starting Question Paper Analysis Script ---")
    
    gemini_model = configure_gemini_api()
    if not gemini_model:
        logging.critical("Exiting due to API configuration failure.")
        # Added a small pause so the user can read the console error
        time.sleep(3)
        return

    # Set up the Tkinter root window for the file dialog
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(
        title="Select the question paper Word file",
        filetypes=(("Word Documents", "*.docx"), ("All files", "*.*"))
    )

    if not file_path:
        logging.info("No file selected. Operation canceled by the user.")
        return

    content = read_docx_content(file_path)
    if not content:
        messagebox.showerror("Error", "Could not read content from the selected Word document.")
        return
        
    unique_content = deduplicate_content(content)
    
    analysis_result_text = analyze_with_gemini(gemini_model, unique_content)
    if not analysis_result_text:
        messagebox.showerror("Error", "Failed to get a valid response from the analysis service.")
        return

    parsed_data = parse_api_response(analysis_result_text)
    if not parsed_data:
        messagebox.showerror("Error", "Could not parse the data from the API response. It may have been empty or malformed.")
        return
        
    create_analysis_document(parsed_data, file_path)
    
    logging.info("--- Process Completed Successfully ---")

if __name__ == '__main__':
    main()