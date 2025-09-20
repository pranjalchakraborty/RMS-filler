import google.generativeai as genai
import docx
import os
import tkinter as tk
from tkinter import filedialog
import logging
import json

# --- Setup basic logging to a file and the console ---
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    handlers=[
                        logging.FileHandler("api_routine_processor.log"),
                        logging.StreamHandler()
                    ])

def select_file(prompt_title):
    """Opens a GUI window for the user to select a .docx file."""
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(
        title=prompt_title,
        filetypes=(("Word Documents", "*.docx"), ("All files", "*.*"))
    )
    return file_path

def select_save_path(prompt_title):
    """Opens a GUI 'Save As' window for the user to choose a save location."""
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.asksaveasfilename(
        title=prompt_title,
        defaultextension=".docx",
        filetypes=(("Word Documents", "*.docx"), ("All files", "*.*"))
    )
    return file_path

def configure_gemini_api():
    """Configures the Gemini API with the user's key."""
    try:
        api_key = os.environ.get("GOOGLE_API_KEY")
        if not api_key:
            logging.warning("GOOGLE_API_KEY environment variable not found.")
            api_key = input("Please enter your Google Gemini API key: ")
        
        genai.configure(api_key=api_key)
        logging.info("Gemini API configured successfully.")
        return genai.GenerativeModel('gemini-2.5-pro')
    except Exception as e:
        logging.error(f"Error configuring Gemini API: {e}")
        return None

# --- NEW: Definitive function to extract a perfect structural blueprint ---
def get_document_context(doc_path):
    """
    Extracts all paragraph text and creates a perfect JSON blueprint of the first table,
    including dimensions and a list of unique cells with their text, width, and spans.
    """
    try:
        doc = docx.Document(doc_path)
        
        # 1. Extract all paragraph text
        doc_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        if not doc.tables:
            logging.error(f"No tables found in {doc_path}")
            return None
        
        table = doc.tables[0]
        
        # 2. Analyze table structure into a JSON blueprint
        total_rows = len(table.rows)
        total_columns = len(table.columns)
        visited_cells = set()
        cell_definitions = []

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell in visited_cells:
                    continue
                
                visited_cells.add(cell)
                
                # Determine the span of this unique cell
                start_row, start_col = r_idx, c_idx
                end_row, end_col = start_row, start_col

                for c in range(c_idx + 1, total_columns):
                    if table.cell(r_idx, c) == cell:
                        end_col = c
                    else:
                        break
                
                for r in range(r_idx + 1, total_rows):
                    if table.cell(r, c_idx) == cell:
                        end_row = r
                    else:
                        break

                cell_definitions.append({
                    "text": cell.text.replace('\n', ' ').strip(),
                    "width": cell.width.emu if cell.width else 0,
                    "start_row": start_row,
                    "start_col": start_col,
                    "end_row": end_row,
                    "end_col": end_col
                })
        
        blueprint = {
            "total_rows": total_rows,
            "total_columns": total_columns,
            "cells": cell_definitions
        }
        
        logging.info(f"Successfully created JSON blueprint for {os.path.basename(doc_path)}")
        return {"document_text": doc_text, "table_blueprint": blueprint}

    except Exception as e:
        logging.error(f"Error extracting document context for {doc_path}: {e}")
        return None

# --- NEW: Master prompt for the API to generate fill instructions ---
def generate_fill_instructions(model, routine1_context, routine2_context, clean_template_context):
    """
    Provides the API with full context and asks it to return a JSON array
    of instructions for filling the clean routine.
    """
    prompt = f"""
    You are an expert system for academic schedule consolidation. Your task is to analyze two source routines and generate a set of precise instructions to fill a clean template for a teacher named "RC".

    **CONTEXT AND DATA BLUEPRINTS:**

    --- CONTEXT FOR ROUTINE 1 ---
    Document Text (contains semester info):
    {routine1_context['document_text']}

    Table Blueprint (a JSON object describing the table structure):
    {json.dumps(routine1_context['table_blueprint'], indent=2)}
    --- END OF CONTEXT FOR ROUTINE 1 ---

    --- CONTEXT FOR ROUTINE 2 ---
    Document Text (contains semester info):
    {routine2_context['document_text']}

    Table Blueprint:
    {json.dumps(routine2_context['table_blueprint'], indent=2)}
    --- END OF CONTEXT FOR ROUTINE 2 ---
    
    --- CLEAN TEMPLATE BLUEPRINT ---
    (This is the target structure. Note its dimensions and cell coordinates.)
    {json.dumps(clean_template_context['table_blueprint'], indent=2)}
    --- END OF CLEAN TEMPLATE BLUEPRINT ---

    **YOUR TASK:**

    Your goal is to generate a JSON array of "fill instructions". Do not generate a table.

    1.  **Identify Semesters:** Read the 'Document Text' for Routine 1 and Routine 2 to determine the semester for each (e.g., "3rd sem", "5th sem").
    2.  **Find RC Classes:** Analyze the 'cells' list in the blueprints for Routine 1 and Routine 2. Find all cell objects where the `text` contains the "RC" marker.
    3.  **Generate Instructions:** For each "RC" class you find, you must generate one or more fill instructions.
        -   Extract the subject name (e.g., "Biomedical Instrumentation").
        -   Append the correct semester tag you found in Step 1 (e.g., "Biomedical Instrumentation (5th sem)").
        -   A single source cell might be merged, spanning multiple coordinates (e.g., from `start_col: 1` to `end_col: 3`). You MUST generate a separate fill instruction object for **each coordinate** it covers.
    4.  **Handle Conflicts:** If two different RC classes are scheduled for the exact same coordinate `(r, c)`, combine their text into a single instruction for that coordinate, separated by a newline (`\\n`).
    5.  **Final Output:** Your response must ONLY be a valid JSON array of instruction objects. Each object must have three keys: `row`, `column`, and `text_to_fill`.

    Example Output Format:
    [
      {{ "row": 1, "column": 1, "text_to_fill": "Biomedical Instrumentation (5th sem)" }},
      {{ "row": 1, "column": 2, "text_to_fill": "Biomedical Instrumentation (5th sem)" }},
      {{ "row": 2, "column": 6, "text_to_fill": "Programming in C (3rd sem)" }}
    ]
    """
    logging.info("--- PROMPT SENT TO GEMINI API ---")
    logging.info(prompt)
    logging.info("---------------------------------")
    
    try:
        response = model.generate_content(prompt)
        cleaned_response = response.text.strip().replace('```json', '').replace('```', '').strip()
        logging.info("--- RAW RESPONSE FROM GEMINI API (Fill Instructions) ---")
        logging.info(cleaned_response)
        logging.info("-------------------------------------------------------")
        return cleaned_response
    except Exception as e:
        logging.error(f"An error occurred during the API call: {e}")
        return None

# --- NEW: Executes the API's instructions ---
def fill_and_save_routine(clean_routine_path, fill_instructions_json_str, save_path):
    """
    Opens the clean routine template and executes the list of fill instructions
    provided by the API.
    """
    try:
        instructions = json.loads(fill_instructions_json_str)
        
        doc = docx.Document(clean_routine_path)
        if not doc.tables:
            logging.error(f"No table found in the clean routine file: {clean_routine_path}")
            return False

        table = doc.tables[0]
        
        # Execute each instruction from the API
        for instruction in instructions:
            r, c = instruction['row'], instruction['column']
            text_to_fill = instruction['text_to_fill']
            
            if r < len(table.rows) and c < len(table.columns):
                cell = table.cell(r, c)
                cell.text = text_to_fill
            else:
                logging.warning(f"Instruction for out-of-bounds cell ({r},{c}) was ignored.")

        doc.save(save_path)
        logging.info(f"Successfully filled and saved the routine to: {save_path}")
        return True

    except json.JSONDecodeError as e:
        logging.error(f"Failed to parse JSON instructions from API: {e}")
        logging.error(f"Invalid JSON received: {fill_instructions_json_str}")
        return False
    except Exception as e:
        logging.error(f"An error occurred while filling and saving the final routine: {e}")
        return False

def main():
    """Main function to drive the script."""
    logging.info("--- Starting API-Powered Routine Processor ---")

    gemini_model = configure_gemini_api()
    if not gemini_model:
        logging.critical("Exiting due to API configuration failure.")
        return

    routine1_path = select_file("Select the FIRST Filled Routine (Source 1)")
    if not routine1_path: return

    routine2_path = select_file("Select the SECOND Filled Routine (Source 2)")
    if not routine2_path: return
    
    clean_routine_path = select_file("Select the Clean Empty Routine (Template)")
    if not clean_routine_path: return

    logging.info("Analyzing files and creating structural blueprints...")
    routine1_context = get_document_context(routine1_path)
    routine2_context = get_document_context(routine2_path)
    clean_template_context = get_document_context(clean_routine_path)
    
    if not all([routine1_context, routine2_context, clean_template_context]):
        logging.critical("Could not create blueprints for one or more DOCX files. Exiting.")
        return
    
    fill_instructions_json_str = generate_fill_instructions(gemini_model, routine1_context, routine2_context, clean_template_context)
    if not fill_instructions_json_str:
        logging.error("Failed to generate fill instructions from the API. Exiting.")
        return
        
    save_path = select_save_path("Save Filled Routine As...")
    if not save_path:
        logging.warning("Operation cancelled. The filled routine was not saved.")
        return
        
    fill_and_save_routine(clean_routine_path, fill_instructions_json_str, save_path)
    
    logging.info("--- Process completed. ---")


if __name__ == "__main__":
    main()