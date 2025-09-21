import pandas as pd
from docx import Document
from docx.shared import Inches
import re
import os
import tkinter as tk
from tkinter import filedialog

def clean_text(text):
    """
    Standardizes a string by removing extra whitespace and newlines.
    This ensures that topics are grouped correctly even with minor formatting differences.
    """
    if not isinstance(text, str):
        return ""
    # Replace newlines and tabs with a space, then reduce multiple spaces to one
    cleaned = re.sub(r'\s+', ' ', text).strip()
    return cleaned

def get_sort_key(topic_string):
    """
    Creates a key for sorting topics.
    - Numbered topics (e.g., "1.1", "1.12") are sorted numerically.
    - Non-numbered topics are placed at the end, sorted alphabetically.
    """
    match = re.match(r'^(\d+)\.(\d+)', topic_string)
    if match:
        # Return a tuple of integers for correct numerical sorting (e.g., 1.10 > 1.2)
        major = int(match.group(1))
        minor = int(match.group(2))
        return (0, major, minor)
    else:
        # Return a tuple that ensures non-numbered topics go last
        return (1, topic_string)

def generate_course_files():
    """
    Main function to orchestrate the process of reading an Excel file and
    generating Word document course files.
    """
    # --- 1. Get Excel file path from user ---
    root = tk.Tk()
    root.withdraw()  # Hide the main tkinter window
    excel_path = filedialog.askopenfilename(
        title="Select the Excel Class Log File",
        filetypes=[("Excel Files", "*.xlsx *.xls")]
    )

    if not excel_path:
        print("No file selected. Exiting program.")
        return

    print(f"Processing file: {excel_path}")

    # --- 2. Read and Prepare Data ---
    try:
        df = pd.read_excel(excel_path)
        
        # --- Clean up column names to prevent KeyErrors ---
        df.columns = df.columns.str.strip() ### <-- NEW/IMPROVED LINE ###
        print(f"Columns found in Excel file: {df.columns.to_list()}") ### <-- NEW/IMPROVED LINE (for debugging) ###

        # Ensure required columns exist after cleaning
        required_cols = ['Class Executed', 'Subject Name', 'Topic Covered']
        if not all(col in df.columns for col in required_cols):
            print(f"Error: The Excel file must contain the columns: {', '.join(required_cols)}")
            print("Please check for typos or missing columns in your Excel file.")
            return

    except Exception as e:
        print(f"Error reading or processing the Excel file: {e}")
        return

    # Filter for Executed classes and drop rows with no subject/topic
    df = df[df['Class Executed'] == 'Yes'].copy()
    df.dropna(subset=['Subject Name', 'Topic Covered'], inplace=True)

    # Clean the 'Topic Covered' column for accurate counting
    df['Cleaned Topic'] = df['Topic Covered'].apply(clean_text)

    # --- 3. Process Each Subject ---
    unique_subjects = df['Subject Name'].unique()

    if len(unique_subjects) == 0:
        print("No subjects found with Executed classes. No files were generated.")
        return

    for subject in unique_subjects:
        print(f"Generating course file for: {subject}...")

        # Filter data for the current subject
        subject_df = df[df['Subject Name'] == subject]

        # Aggregate topics and count occurrences (hours)
        topic_counts = subject_df['Cleaned Topic'].value_counts().to_dict()

        # Sort topics using the custom sort key
        sorted_topics = sorted(topic_counts.keys(), key=get_sort_key)

        # --- 4. Generate the Word Document ---
        doc = Document()
        doc.add_heading(f'Course File: {subject}', level=1)
        doc.add_paragraph() # Add some space

        # Create the table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False # Allow manual column sizing
        
        # Set column widths
        table.columns[0].width = Inches(0.5) # Sl. No.
        table.columns[1].width = Inches(4.0) # Topics Covered
        table.columns[2].width = Inches(1.0) # Methodology
        table.columns[3].width = Inches(0.5) # Hours
        table.columns[4].width = Inches(0.8) # Remarks

        # Populate header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sl. No.'
        hdr_cells[1].text = 'Topics Covered'
        hdr_cells[2].text = 'Methodology Used'
        hdr_cells[3].text = 'Hours'
        hdr_cells[4].text = 'Remarks'

        # Populate data rows
        for i, topic in enumerate(sorted_topics, 1):
            hours = topic_counts[topic]
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = topic
            row_cells[2].text = 'Chalk & Talk'
            row_cells[3].text = f"{hours:02}" # Format hours as 01, 02, etc.
            row_cells[4].text = '' # Remarks field is empty

        # --- 5. Save the Document ---
        # Sanitize the subject name for use in a filename
        safe_subject_name = re.sub(r'[\\/*?:"<>|]', "", subject)
        output_filename = f"CourseFile_{safe_subject_name}.docx"

        try:
            doc.save(output_filename)
            print(f"Successfully saved '{output_filename}'")
        except Exception as e:
            print(f"Error saving file for {subject}: {e}")

# --- Run the main function ---
if __name__ == "__main__":
    generate_course_files()