import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import time
import os
import sys
import google.generativeai as genai
import logging
import pandas as pd
import re
import json
from copy import deepcopy

# --- Setup basic logging to a file and the console ---
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    handlers=[
                        logging.FileHandler("integrated_course_file_generator.log"),
                        logging.StreamHandler()
                    ])

# ==============================================================================
# SECTION 1: GLOBAL API CONFIGURATION AND HELPER FUNCTIONS
# ==============================================================================

GEMINI_MODEL = None

def get_gemini_model():
    """Configures and returns a Gemini model instance, requesting API key only once."""
    global GEMINI_MODEL
    if GEMINI_MODEL is not None:
        return GEMINI_MODEL
    try:
        api_key = os.environ.get("GOOGLE_API_KEY")
        if not api_key:
            api_key = input("Please enter your Google Gemini API key once and press Enter: ")
        if not api_key:
            messagebox.showerror("API Error", "API Key was not provided.")
            return None
        genai.configure(api_key=api_key)
        logging.info("Gemini API configured successfully for the session.")
        # FINAL FIX: Using the correct, more capable model as per the proven script
        GEMINI_MODEL = genai.GenerativeModel('gemini-2.5-pro')
        return GEMINI_MODEL
    except Exception as e:
        messagebox.showerror("API Error", f"Failed to configure the Gemini API: {e}")
        return None

def select_file(prompt_title, file_types=(("Word Documents", "*.docx"), ("All files", "*.*"))):
    root = tk.Tk(); root.withdraw()
    return filedialog.askopenfilename(title=prompt_title, filetypes=file_types)

def find_paragraph(doc, text_to_find):
    for p in doc.paragraphs:
        if text_to_find in p.text: return p
    logging.warning(f"Could not find paragraph with text: '{text_to_find}'")
    return None

def paste_text_after_paragraph(text_content, anchor_paragraph, main_doc):
    insert_element = anchor_paragraph._p
    if not text_content.strip(): return anchor_paragraph
    for line in text_content.splitlines():
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r'); new_t = OxmlElement('w:t')
        new_t.text = line
        new_r.append(new_t); new_p.append(new_r)
        insert_element.addnext(new_p)
        insert_element = new_p
    logging.info("Successfully pasted text content.")
    return Paragraph(insert_element, main_doc._body)

def _remove_width_specifications(tbl_xml):
    """Removes all hard-coded width settings from a table's XML."""
    for tblW in tbl_xml.xpath('.//w:tblW'): tblW.getparent().remove(tblW)
    for tcW in tbl_xml.xpath('.//w:tcW'): tcW.getparent().remove(tcW)

def set_table_autofit(table):
    """Sets a table to autofit to the window width using robust XML manipulation."""
    try:
        tblPr = table._element.xpath('w:tblPr')
        if not tblPr:
            tblPr = OxmlElement('w:tblPr'); table._element.insert(0, tblPr)
        else:
            tblPr = tblPr[0]
        tblAutofit = OxmlElement('w:tblAutofit')
        tblAutofit.set(qn('w:val'), 'autofitWindow')
        tblPr.append(tblAutofit)
    except Exception as e:
        logging.error(f"Could not set table to autofit: {e}")

def paste_table_after_paragraph(source_table, anchor_paragraph, main_doc, add_blank_line_before=False):
    """Copies a table, sanitizes it, pastes it, and applies autofit."""
    try:
        copied_table_xml = deepcopy(source_table._tbl)
        _remove_width_specifications(copied_table_xml)
        
        insert_point = anchor_paragraph._p
        if add_blank_line_before:
            blank_p = OxmlElement('w:p')
            insert_point.addnext(blank_p)
            insert_point = blank_p
        
        insert_point.addnext(copied_table_xml)
        new_table_element = insert_point.getnext()

        for table in main_doc.tables:
            if table._element is new_table_element:
                set_table_autofit(table)
                break
        logging.info("Successfully pasted and autofit a table.")
        return True
    except Exception as e:
        logging.error(f"Error while pasting table: {e}")
        return False

def extract_course_details(course_file_path):
    try:
        doc = Document(course_file_path)
        course_name, semester = None, None
        for para in doc.paragraphs:
            if course_name is None:
                match = re.search(r"Course:\s*(.*)", para.text, re.IGNORECASE)
                if match: course_name = match.group(1).strip()
            if semester is None and "Semester:" in para.text:
                semester = para.text.split(":")[-1].strip()
            if course_name and semester: break
        if course_name and semester:
            logging.info(f"Extracted Course: '{course_name}', Semester: '{semester}'")
            return course_name, semester
        return None, None
    except Exception as e:
        logging.error(f"Failed to read course file: {e}")
        return None, None

def normalize_string(s):
    return re.sub(r'[^a-z0-9]', '', str(s).lower())

def get_script_directory():
    """Returns the directory of the script or the compiled executable."""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))

# ==============================================================================
# SECTION 2: BLOOMS TAXONOMY
# ==============================================================================

def run_blooms_process(main_course_doc):
    logging.info("\n--- Starting Question Paper Analysis and Integration ---")
    gemini_model = get_gemini_model()
    if not gemini_model: return

    for i in range(1, 3):
        internal_num = "first" if i == 1 else "second"
        qp_path = select_file(f"Select the {internal_num} internal question paper")
        if not qp_path: continue
        
        logging.info(f"Processing Internal {i} paper: {os.path.basename(qp_path)}")

        try:
            content = "\n".join([p.text for p in Document(qp_path).paragraphs if p.text.strip()])
            unique_content = "\n".join(list(dict.fromkeys(line.strip() for line in content.splitlines() if line.strip())))
            
            prompt = f"""
            Analyze the following question paper based on Bloom's Taxonomy. 
            Provide the output ONLY in a pipe-separated format without headers.
            Each line in your output must represent a single, specific question subdivision (e.g., 1a, 1i, 2b), NOT the aggregated question number (e.g., 1, 2).
            Format: Question_No|Marks|Blooms_Taxonomy_Level

            Question Paper Text:
            ---
            {unique_content}
            ---
            """
            analysis_result_text = gemini_model.generate_content(prompt).text
            
            parsed_data = []
            for line in analysis_result_text.strip().split('\n'):
                parts = [part.strip() for part in line.split('|')]
                if len(parts) == 3: parsed_data.append(parts)
            
            if not parsed_data: continue

            analysis_doc = Document()
            table = analysis_doc.add_table(rows=1, cols=3, style='Table Grid')
            set_table_autofit(table)
            hdr = ['Question No.', 'Marks', 'Blooms Taxonomy Level']
            for idx, text in enumerate(hdr): table.cell(0, idx).text = text
            for item in parsed_data:
                cells = table.add_row().cells
                cells[0].text = item[0]; cells[1].text = item[1]; cells[2].text = item[2]
            
            qp_anchor = find_paragraph(main_course_doc, f"Question Paper: Internal {i}")
            if qp_anchor:
                text_anchor = paste_text_after_paragraph(unique_content, qp_anchor, main_course_doc)
                paste_table_after_paragraph(table, text_anchor, main_course_doc, add_blank_line_before=True)
            
        except Exception as e:
            logging.error(f"An error occurred during Blooms process for Internal {i}: {e}")

# ==============================================================================
# SECTION 3: ROUTINE GENERATION
# ==============================================================================

def extract_semester_from_text(text):
    match = re.search(r'(\d+)(?:st|nd|rd|th)?\s+semester', text, re.IGNORECASE)
    if match: return match.group(1)
    return None

def routine_get_document_context(doc_path):
    """The proven, correct blueprint generation logic from the reference script."""
    try:
        doc = Document(doc_path)
        doc_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        if not doc.tables:
            return None
        
        table = doc.tables[0]
        total_rows, total_columns = len(table.rows), len(table.columns)
        visited_cells, cell_definitions = set(), []

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell in visited_cells:
                    continue
                visited_cells.add(cell)
                start_row, start_col, end_row, end_col = r_idx, c_idx, r_idx, c_idx

                for c in range(c_idx + 1, total_columns):
                    if table.cell(r_idx, c) == cell:
                        end_col = c
                    else:
                        break
                for r in range(r_idx + 1, total_rows):
                    if table.cell(r, c_idx) == cell:
                        end_row = r
                    else:
                        break
                cell_definitions.append({
                    "text": cell.text.replace('\n', ' ').strip(),
                    "width": cell.width.emu if cell.width else 0,
                    "start_row": start_row, "start_col": start_col,
                    "end_row": end_row, "end_col": end_col
                })
        
        blueprint = {
            "total_rows": total_rows,
            "total_columns": total_columns,
            "cells": cell_definitions
        }
        logging.info(f"Successfully created JSON blueprint for {os.path.basename(doc_path)}")
        return {"document_text": doc_text, "table_blueprint": blueprint, "table": table, "path": doc_path}
    except Exception as e:
        logging.error(f"Error extracting document context for {doc_path}: {e}")
        return None

def routine_generate_fill_instructions(model, r1_ctx, r2_ctx, cl_ctx, initials):
    """The proven, correct prompt from the reference script."""
    prompt = f"""
    You are an expert system for academic schedule consolidation. Your task is to analyze two source routines and generate a set of precise instructions to fill a clean template for a teacher with the initials "{initials}".

    **CONTEXT AND DATA BLUEPRINTS:**
    --- CONTEXT FOR ROUTINE 1 ---
    Document Text (contains semester info):
    {r1_ctx['document_text']}
    Table Blueprint (a JSON object describing the table structure):
    {json.dumps(r1_ctx['table_blueprint'], indent=2)}
    --- END OF CONTEXT FOR ROUTINE 1 ---
    --- CONTEXT FOR ROUTINE 2 ---
    Document Text (contains semester info):
    {r2_ctx['document_text']}
    Table Blueprint:
    {json.dumps(r2_ctx['table_blueprint'], indent=2)}
    --- END OF CONTEXT FOR ROUTINE 2 ---
    --- CLEAN TEMPLATE BLUEPRINT ---
    (This is the target structure. Note its dimensions and cell coordinates.)
    {json.dumps(cl_ctx['table_blueprint'], indent=2)}
    --- END OF CLEAN TEMPLATE BLUEPRINT ---

    **YOUR TASK:**
    Your goal is to generate a JSON array of "fill instructions". Do not generate a table.
    1.  **Identify Semesters:** Read the 'Document Text' for Routine 1 and Routine 2 to determine the semester for each.
    2.  **Find Teacher's Classes:** Analyze the 'cells' list in the blueprints for Routine 1 and Routine 2. Find all cell objects where the `text` contains the marker "{initials}".
    3.  **Generate Instructions:** For each class you find, generate one or more fill instructions. Append the correct semester tag.
    4.  A single source cell might be merged. You MUST generate a separate fill instruction object for **each coordinate** it covers.
    5.  **Handle Conflicts:** If two different classes for "{initials}" are scheduled for the exact same coordinate `(r, c)`, combine their text into a single instruction, separated by a newline (`\\n`).
    6.  **Final Output:** Your response must ONLY be a valid JSON array of instruction objects. Each object must have three keys: `row`, `column`, and `text_to_fill`.
    """
    try:
        response = model.generate_content(prompt)
        return response.text.strip().replace('```json', '').replace('```', '').strip()
    except Exception as e:
        logging.error(f"API call for routine instructions failed: {e}")
        return None

def routine_fill_and_save_routine(clean_path, instructions_json, save_path):
    """Fills a clean template using API instructions."""
    try:
        instructions = json.loads(instructions_json)
        doc = Document(clean_path)
        if not doc.tables: return False
        table = doc.tables[0]
        for inst in instructions:
            if 'row' in inst and 'column' in inst and 'text_to_fill' in inst:
                r, c, txt = inst['row'], inst['column'], inst['text_to_fill']
                if r < len(table.rows) and c < len(table.columns):
                    table.cell(r, c).text = txt
        doc.save(save_path)
        logging.info(f"Successfully generated and saved routine to: {save_path}")
        return True
    except (json.JSONDecodeError, KeyError) as e:
        logging.error(f"Error filling/saving routine: {e}")
        return False

def run_routine_process(main_course_doc, semester_from_course):
    """Main logic to orchestrate routine generation and pasting."""
    logging.info("\n--- Starting Routine Generation and Integration ---")
    teacher_initials = "RC"
    script_dir = get_script_directory()
    personal_routine_path = os.path.join(script_dir, f"personal_routine_{teacher_initials}.docx")
    
    r1_path = select_file("Select FIRST Class Routine (Source 1)")
    r2_path = select_file("Select SECOND Class Routine (Source 2)")
    if not (r1_path and r2_path):
        logging.warning("Both class routines not selected. Skipping routine processing.")
        return

    r1_ctx = routine_get_document_context(r1_path)
    r2_ctx = routine_get_document_context(r2_path)

    run_generation = False
    if os.path.exists(personal_routine_path):
        root = tk.Tk(); root.withdraw()
        if messagebox.askyesno(title="File Exists", message=f"'{os.path.basename(personal_routine_path)}' already exists.\nRegenerate it?"):
            run_generation = True
    else:
        run_generation = True

    if run_generation:
        logging.info("Starting personal routine generation process...")
        gemini_model = get_gemini_model()
        if not gemini_model: return

        clean_path = os.path.join(script_dir, "clean_routine.docx")
        if not os.path.exists(clean_path):
            logging.error(f"'clean_routine.docx' not found in the script's directory: {script_dir}")
            messagebox.showerror("File Not Found", f"'clean_routine.docx' must be in the same folder as the script/exe.")
            return
        
        logging.info(f"Using clean routine template: {clean_path}")
        
        cl_ctx = routine_get_document_context(clean_path)
        if not (r1_ctx and r2_ctx and cl_ctx):
            logging.error("Could not process one or more routine files. Aborting generation.")
            return

        instructions = routine_generate_fill_instructions(gemini_model, r1_ctx, r2_ctx, cl_ctx, teacher_initials)
        if instructions:
            routine_fill_and_save_routine(clean_path, instructions, personal_routine_path)
        else:
            logging.error("Failed to get generation instructions from API.")

    course_sem_num = re.search(r'\d+', semester_from_course).group()
    correct_ctx = None
    if r1_ctx and extract_semester_from_text(r1_ctx['document_text']) == course_sem_num: correct_ctx = r1_ctx
    elif r2_ctx and extract_semester_from_text(r2_ctx['document_text']) == course_sem_num: correct_ctx = r2_ctx
    
    class_anchor = find_paragraph(main_course_doc, "Class Timetable:")
    if class_anchor and correct_ctx:
        paste_table_after_paragraph(correct_ctx['table'], class_anchor, main_course_doc)

    if os.path.exists(personal_routine_path):
        individual_anchor = find_paragraph(main_course_doc, "Individual Timetable:")
        if individual_anchor:
            personal_doc = Document(personal_routine_path)
            if personal_doc.tables:
                paste_table_after_paragraph(personal_doc.tables[0], individual_anchor, main_course_doc)
    else:
        logging.warning("Personal routine file not found for pasting.")

# ==============================================================================
# SECTION 4: COURSE TABLE GENERATION
# ==============================================================================

def generate_course_files():
    """Generates course files from a user-selected Excel file."""
    # FINAL FIX: Restore user prompt for selecting the Excel file.
    excel_path = select_file("Select the Excel Class Log File", [("Excel Files", "*.xlsx *.xls")])
    if not excel_path:
        logging.warning("No Excel file selected by user. Aborting course table generation.")
        return
    
    script_dir = get_script_directory()
    try:
        df = pd.read_excel(excel_path, engine='openpyxl')
        df.columns = df.columns.str.strip()
        required = ['Class Executed', 'Subject Name', 'Topic Covered']
        if not all(col in df.columns for col in required): return

        df = df[df['Class Executed'] == 'Yes'].dropna(subset=['Subject Name', 'Topic Covered']).copy()
        
        for subject in df['Subject Name'].unique():
            doc = Document()
            table = doc.add_table(rows=1, cols=5, style='Table Grid')
            set_table_autofit(table)
            
            hdr = ['Sl. No.', 'Topics Covered', 'Methodology Used', 'Hours', 'Remarks']
            for i, h in enumerate(hdr): table.rows[0].cells[i].text = h
            
            subject_df = df[df['Subject Name'] == subject]
            topic_counts = subject_df['Topic Covered'].value_counts()
            
            for i, (topic, hours) in enumerate(topic_counts.items(), 1):
                cells = table.add_row().cells
                cells[0].text = str(i); cells[1].text = str(topic)
                cells[2].text = 'Chalk & Talk'; cells[3].text = f"{int(hours):02}"
            
            safe_subject = re.sub(r'[\\/*?:"<>|]', "", str(subject))
            save_path = os.path.join(script_dir, f"CourseFile_{safe_subject}.docx")
            doc.save(save_path)
            logging.info(f"Successfully saved '{os.path.basename(save_path)}'")
    except Exception as e:
        logging.error(f"Error during Excel processing: {e}")

def run_course_table_process(main_course_doc, course_name):
    logging.info("\n--- Starting Course Table Generation and Integration ---")
    script_dir = get_script_directory()
    normalized_course_name = normalize_string(course_name)
    
    safe_subject_name = re.sub(r'[\\/*?:"<>|]', "", course_name)
    target_filename = f"CourseFile_{safe_subject_name}.docx"
    logging.info(f"Searching for course table file: '{target_filename}'")

    found_file_path = None
    for filename in os.listdir(script_dir):
        if filename.lower().startswith('coursefile_') and filename.lower().endswith('.docx'):
            if normalize_string(filename[11:-5]) == normalized_course_name:
                found_file_path = os.path.join(script_dir, filename)
                logging.info(f"Match found. Using file: '{filename}'")
                break
    
    regenerate = False
    if found_file_path:
        root = tk.Tk(); root.withdraw()
        if messagebox.askyesno("File Exists", f"'{os.path.basename(found_file_path)}' already exists.\nRegenerate it from Excel?"):
            regenerate = True
    else:
        regenerate = True

    if regenerate:
        logging.info("Running course table generation from Excel.")
        generate_course_files()
        found_file_path = None # Reset to search again
        for filename in os.listdir(script_dir):
            if filename.lower().startswith('coursefile_') and filename.lower().endswith('.docx'):
                if normalize_string(filename[11:-5]) == normalized_course_name:
                    found_file_path = os.path.join(script_dir, filename)
                    break

    if found_file_path:
        anchor = find_paragraph(main_course_doc, "Lecture schedule with methodology being used/adopted")
        if anchor:
            course_table_doc = Document(found_file_path)
            if course_table_doc.tables:
                paste_table_after_paragraph(course_table_doc.tables[0], anchor, main_course_doc)
    else:
        logging.error(f"Could not find or generate course table for '{course_name}'.")

# ==============================================================================
# SECTION 5: MAIN EXECUTION BLOCK
# ==============================================================================

def main():
    logging.info("--- Starting Integrated Course File Generator ---")
    course_file_path = select_file("Select the Main Course File to be Populated")
    if not course_file_path: return

    course_name, semester = extract_course_details(course_file_path)
    if not (course_name and semester):
        messagebox.showerror("Error", "Could not extract Course Name/Semester.")
        return
        
    main_doc = Document(course_file_path)

    run_blooms_process(main_doc)
    run_routine_process(main_doc, semester)
    run_course_table_process(main_doc, course_name)

    try:
        root = tk.Tk(); root.withdraw()
        if messagebox.askyesno("Save Action", "Do you want to OVERWRITE the original file?\n\n(Choosing 'No' will create a new 'INTEGRATED_...' file instead.)"):
            save_path = course_file_path
            logging.info(f"User chose to overwrite original file.")
        else:
            save_path = os.path.join(os.path.dirname(course_file_path), f"INTEGRATED_{os.path.basename(course_file_path)}")
            logging.info(f"User chose to create a new file.")
        
        main_doc.save(save_path)
        logging.info(f"--- Process Completed. Final document saved as: {save_path} ---")
        messagebox.showinfo("Success", f"All tasks completed!\nIntegrated file saved as:\n{save_path}")
    except Exception as e:
        messagebox.showerror("Save Error", f"Could not save final document.\nError: {e}")

if __name__ == "__main__":
    main()